# notebook-backend/app.py

from flask import Flask, request, jsonify, send_from_directory, send_file
# 确保这里导入了 cross_origin
from flask_cors import CORS, cross_origin # <--- 确保这一行正确导入
import os
from werkzeug.utils import secure_filename
import json
import time
from datetime import datetime
from dotenv import load_dotenv
import re
import subprocess
from pathlib import Path
import threading
import tempfile # <--- 确保这一行正确导入
import shutil # <--- 确保这一行正确导入
import sys

# DeepSeek 相关的导入
from openai import OpenAI
from openai import APIConnectionError, RateLimitError, APIStatusError

# 用于文档内容提取和生成
import docx
from io import BytesIO
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter

# 导入中文分词库 jieba
import jieba

# 导入业务逻辑模块
from llm_interface import DeepSeekLLM
from data_manager import TemplateMethodManager
from text_summarizer import TextSummarizer
from question_rewriter import QuestionRewriter

# 加载 .env 文件中的环境变量
load_dotenv()

# --- Flask 应用配置 ---
app = Flask(__name__)

# --- CRITICAL FIX: CORS Configuration ---
CORS(app, resources={
    r"/*": {"origins": "http://localhost:3000", "methods": ["GET", "HEAD", "POST", "PUT", "DELETE", "PATCH", "OPTIONS"],
           "allow_headers": ["Content-Type", "x-requested-with", "Authorization"], "supports_credentials": True}})

print("DEBUG: CORS configured globally for all routes in app.py")

# --- 文件存储配置 ---
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
    print(f"Created uploads directory at: {os.path.abspath(UPLOAD_FOLDER)}")

# --- API Keys 配置 ---
deepseek_api_key = os.getenv("DEEPSEEK_API_KEY", "sk-945bab02da964bcdaca673485c32dfff") # Added default from user code
if not deepseek_api_key:
    print("WARNING: DEEPSEEK_API_KEY not found in .env file or environment variables.")
    print("AI model calls might fail. Please add DEEPSEEK_API_KEY=your_key_here to your .env file.")

dashscope_api_key = os.getenv("DASHSCOPE_API_KEY", "sk-31510d140d0c4af28612ce447f28943e") # Added default from user code
if not dashscope_api_key:
    print("WARNING: DASHSCOPE_API_KEY not found in .env file or environment variables.")
    print("External processing scripts (OCR, Catalog, Segmentation) might fail.")

DEEPSEEK_BASE_URL = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1")
openai_client_for_chat_rag = OpenAI(
    api_key=deepseek_api_key,
    base_url=DEEPSEEK_BASE_URL
)
DEFAULT_DEEPSEEK_MODEL = os.getenv("DEFAULT_DEEPSEEK_MODEL", "deepseek-chat")


# --- 新功能模块初始化 ---
llm_client_for_tasks = DeepSeekLLM(api_key=deepseek_api_key, model=DEFAULT_DEEPSEEK_MODEL, base_url=DEEPSEEK_BASE_URL)
TEMPLATES_FILE = "templates_methods.json"
template_manager = TemplateMethodManager(file_path=TEMPLATES_FILE)
text_summarizer_module = TextSummarizer(llm_client=llm_client_for_tasks)
question_rewriter_module = QuestionRewriter(llm_client=llm_client_for_tasks, template_manager=template_manager)

# --- 模拟数据持久化 ---
mock_files_metadata = []
mock_chat_history = []
mock_chat_messages = {}

MOCK_DATA_FILE = 'mock_data.json'


def load_mock_data():
    global mock_files_metadata, mock_chat_history, mock_chat_messages
    if os.path.exists(MOCK_DATA_FILE):
        with open(MOCK_DATA_FILE, 'r', encoding='utf-8') as f:
            try:
                data = json.load(f)
                mock_files_metadata = data.get('files', [])
                mock_chat_history = data.get('chat_history', [])
                mock_chat_messages = data.get('chat_messages', {})
                print(f"Mock data loaded from {MOCK_DATA_FILE}")
            except json.JSONDecodeError as e:
                print(f"Error decoding {MOCK_DATA_FILE}: {e}. Starting with empty data.")
                if os.path.exists(MOCK_DATA_FILE):
                    os.rename(MOCK_DATA_FILE, f"{MOCK_DATA_FILE}.bak_{int(time.time())}")
    else:
        print(f"No {MOCK_DATA_FILE} found, starting with empty data.")


def save_mock_data():
    data = {
        'files': mock_files_metadata,
        'chat_history': mock_chat_history,
        'chat_messages': mock_chat_messages
    }
    with open(MOCK_DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4, ensure_ascii=False)
    print(f"Mock data saved to {MOCK_DATA_FILE}")


# --- 文件内容分块函数 (for RAG) ---
def chunktext(text, chunk_size=800, chunk_overlap=100):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        is_separator_regex=False
    )
    raw_chunks = text_splitter.split_text(text)
    return [c.strip() for c in raw_chunks if c.strip()]


# --- 文件内容提取函数 (for RAG and Template Extraction from full text) ---
def extract_text_from_file(filepath, file_type):
    full_text = ""
    try:
        if file_type == 'docx':
            doc = docx.Document(filepath)
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            print(f"Extracted {len(full_text)} characters from DOCX.")
        elif file_type == 'pdf':
            reader = PdfReader(filepath)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    full_text += page_text + "\n"
            print(f"Extracted {len(full_text)} characters from PDF using PyPDF2.")
        elif file_type in ['txt', 'md', 'json', 'log', 'py', 'js', 'css', 'html', 'xml']:
            with open(filepath, 'r', encoding='utf-8') as f:
                full_text = f.read()
            print(f"Read {len(full_text)} characters from plain text file.")
        else:
            print(f"Warning: File type {file_type} not supported for text extraction.")
            return [], f"无法提取类型为 {file_type} 的文件内容。"

    except Exception as e:
        print(f"Error extracting text from {filepath} ({file_type}): {e}")
        return [], f"提取文件内容时发生错误: {str(e)}"

    chunks = chunktext(full_text)

    structured_chunks = []
    for i, chunk_text in enumerate(chunks):
        if chunk_text.strip():
            structured_chunks.append({
                "id": f"chunk_{i + 1}",
                "text": chunk_text,
                "original_full_text": full_text # Storing full text here might be memory intensive for large files
            })
    return structured_chunks, ""


# --- 外部脚本执行函数 ---
def run_processing_pipeline_in_thread(textbook_base_name_unique: str, original_filename_for_subscript: str, processed_dir_full_name: str):
    """
    在后台线程中运行一系列 Python 脚本来处理上传的文件。
    Args:
        textbook_base_name_unique (str): 唯一的教材基础名 (例如 "Essay", "Essay_1"). This is the name WITHOUT _dir.
        original_filename_for_subscript (str): 原始上传的文件名 (例如 "Essay.pdf").
        processed_dir_full_name (str): 完整的处理目录名 (例如 "Essay_dir", "Essay_1_dir"). This is passed to scripts that need the full dir name.
    """
    print(
        f"\n--- Thread starting background processing pipeline for '{textbook_base_name_unique}' (Original: '{original_filename_for_subscript}', Processed Dir: '{processed_dir_full_name}') ---")
    current_script_dir = Path(__file__).parent.resolve()

    # Scripts now might need different arguments.
    # (script_name, arg1_for_script, arg2_for_script, ...)
    # Some scripts might use textbook_base_name_unique (e.g. "MyBook")
    # Others might use processed_dir_full_name (e.g. "MyBook_dir" or "MyBook_Exam_dir")
    scripts_to_run = [
        ("images_and_ocr.py", textbook_base_name_unique, original_filename_for_subscript),
        ("get_catalog.py", textbook_base_name_unique, original_filename_for_subscript),
        ("get_segment.py", textbook_base_name_unique, original_filename_for_subscript),
        ("get_orgchart.py", textbook_base_name_unique, original_filename_for_subscript), # Uses textbook_base_name_unique
        # If get_questions_orgchart.py is part of the standard pipeline:
        # It needs the 'processed_dir_full_name' to know where to look for 'questions/' and where to save its output.
        # ("get_questions_orgchart.py", processed_dir_full_name) 
    ]
    # Note: The decision to run get_questions_orgchart.py might be conditional
    # or triggered by a separate API call if it's not always needed on initial upload.
    # For now, I'm commenting it out from the main pipeline. If you want it here, uncomment and ensure
    # get_questions_orgchart.py handles its arguments correctly.

    env = os.environ.copy()
    if dashscope_api_key:
        env['DASHSCOPE_API_KEY'] = dashscope_api_key
    else:
        print("WARNING: DASHSCOPE_API_KEY not set in environment for sub-processes. Scripts requiring it might fail.")
        # Decide if you want to return or let some scripts run without it

    python_executable = sys.executable

    for script_config in scripts_to_run:
        script_name = script_config[0]
        args_for_script = script_config[1:]

        script_path = current_script_dir / script_name
        if not script_path.exists():
            print(f"WARNING: Script '{script_name}' not found at {script_path}, skipping.")
            continue

        command_args_str = [str(arg) for arg in args_for_script]
        command = [python_executable, str(script_path)] + command_args_str

        print(f"Executing background command: {' '.join(command)}")

        try:
            process = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding='utf-8', # Changed to utf-8 for broader compatibility
                env=env,
                cwd=current_script_dir, # Run script from its own directory
                errors='replace'
            )

            stdout_content = process.stdout
            stderr_content = process.stderr

            if process.returncode == 0:
                print(f"Script '{script_name}' for '{textbook_base_name_unique}' completed successfully.")
                if stdout_content.strip():
                    print(f"  Stdout:\n{stdout_content}")
                if stderr_content.strip(): # Log stderr even on success, as it might contain warnings
                    print(f"  Stderr (Warnings/Info):\n{stderr_content}")
            else:
                print(
                    f"ERROR: Script '{script_name}' for '{textbook_base_name_unique}' failed with exit code {process.returncode}.")
                print(f"  Stdout:\n{stdout_content}")
                print(f"  Stderr:\n{stderr_content}")

        except Exception as e:
            print(f"ERROR: Exception while running script '{script_name}' for '{textbook_base_name_unique}': {e}")

    print(f"--- Background processing pipeline for '{textbook_base_name_unique}' (Dir: '{processed_dir_full_name}') finished. ---")


# --- API 路由 ---

# 1. 提供静态文件 (上传的文件)
@app.route('/uploads/<path:filename_with_path>')
@cross_origin()
def uploaded_file_endpoint(filename_with_path):
    print(f"\n--- DEBUG: Serving file request ---")
    print(f"Requested filename (from URL): {filename_with_path}")

    # This logic correctly serves files from subdirectories like 'Essay_dir/images_dir/page.jpg'
    # or 'Essay_dir/Essay.pdf'
    # The `send_from_directory` first argument is the base directory *within* UPLOAD_FOLDER
    # and the second is the path relative to that base.

    # Example: /uploads/Essay_dir/images_dir/page.jpg
    # UPLOAD_FOLDER = 'uploads'
    # filename_with_path = 'Essay_dir/images_dir/page.jpg'
    # directory_name = 'Essay_dir' (first part of filename_with_path)
    # path_to_file_in_dir = 'images_dir/page.jpg' (rest of filename_with_path)
    # send_from_directory(Path(UPLOAD_FOLDER) / directory_name, path_to_file_in_dir)
    # This resolves to send_from_directory('uploads/Essay_dir', 'images_dir/page.jpg')

    # Example: /uploads/Essay_dir/Essay.pdf
    # directory_name = 'Essay_dir'
    # path_to_file_in_dir = 'Essay.pdf'
    # send_from_directory(Path(UPLOAD_FOLDER) / directory_name, path_to_file_in_dir)
    # This resolves to send_from_directory('uploads/Essay_dir', 'Essay.pdf')

    # It seems your original logic for splitting path_parts was trying to achieve this.
    # A simpler way if `filename_with_path` always contains the sub-directory as the first component:
    
    try:
        # Path objects make manipulation easier
        requested_path = Path(filename_with_path)
        
        # The first part of the path is assumed to be the unique directory (e.g., "Essay_dir")
        # The rest is the path within that directory.
        if not requested_path.parts:
            return jsonify({'error': 'Invalid file path.'}), 400

        # directory_under_uploads is e.g., "Essay_dir"
        directory_under_uploads = requested_path.parts[0]
        # relative_path_within_specific_dir is e.g., "Essay.pdf" or "images_dir/page.jpg"
        relative_path_within_specific_dir = Path(*requested_path.parts[1:])

        # The base directory for send_from_directory is the specific book's folder
        # e.g., 'uploads/Essay_dir'
        base_serving_directory = Path(UPLOAD_FOLDER) / directory_under_uploads
        
        # Check if the base_serving_directory itself exists and is a directory
        if not base_serving_directory.is_dir():
            print(f"ERROR: Base serving directory '{base_serving_directory}' not found or not a directory.")
            return jsonify({'error': 'File resource directory not found.'}), 404

        # The actual file to serve
        full_path_to_serve = base_serving_directory / relative_path_within_specific_dir

        if full_path_to_serve.exists() and full_path_to_serve.is_file():
            print(f"Attempting to serve: {full_path_to_serve}")
            # send_from_directory's first arg is the *absolute or relative path to the directory*,
            # and the second is the *filename relative to that directory*.
            return send_from_directory(str(base_serving_directory.resolve()), str(relative_path_within_specific_dir))
        else:
            # Fallback for older structure or direct file requests in UPLOAD_FOLDER root (less likely with new structure)
            full_path_root = Path(UPLOAD_FOLDER) / filename_with_path
            if full_path_root.exists() and full_path_root.is_file():
                print(f"Attempting to serve from root uploads (fallback): {full_path_root}")
                return send_from_directory(UPLOAD_FOLDER, filename_with_path)
            
            print(f"ERROR: File '{filename_with_path}' not found at expected location: {full_path_to_serve}")
            return jsonify({'error': 'File not found on server.', 'detail': f'File {filename_with_path} does not exist.'}), 404

    except Exception as e:
        print(f"ERROR serving file {filename_with_path}: {e}")
        return jsonify({'error': 'Server error while trying to serve file.'}), 500


# 2. 获取文件列表 (显示原始PDF和其处理状态)
@app.route('/api/files', methods=['GET'])
@cross_origin()
def get_files():
    files_on_disk = []
    # Iterate through items in the UPLOAD_FOLDER
    for item_name in os.listdir(UPLOAD_FOLDER):
        item_path = Path(UPLOAD_FOLDER) / item_name
        # We are looking for directories that end with '_dir' as these are our processed book/item folders
        if item_path.is_dir() and item_name.endswith('_dir'):
            original_file_found = None
            # Try to find a PDF, DOCX, or TXT file inside this directory to represent the "original"
            # This order can be adjusted based on preference.
            possible_extensions = ['.pdf', '.docx', '.txt'] # Add other primary file types if needed
            
            for f_in_dir in item_path.iterdir():
                if f_in_dir.is_file() and f_in_dir.suffix.lower() in possible_extensions:
                    original_file_found = f_in_dir
                    break # Found a primary file

            if original_file_found:
                # Check for a processing marker, e.g., 'catalog_with_segments.json'
                # or 'questions_filtered_textbook_orgchart.json' if that's a final step
                is_processed_flag = (item_path / "textbook_information" / "catalog_with_segments.json").exists()
                # You might want a more comprehensive check for "is_processed"

                files_on_disk.append({
                    'id': item_name,  # Use the directory name (e.g., "Essay_dir") as the unique ID
                    'name': original_file_found.name, # Display original filename (e.g., "Essay.pdf")
                    'size': f"{round(original_file_found.stat().st_size / (1024 * 1024), 2)}MB",
                    'type': original_file_found.suffix.lower().replace('.', ''),
                    'uploadDate': datetime.fromtimestamp(original_file_found.stat().st_mtime).strftime('%Y-%m-%d %H:%M'),
                    'file_type_tag': 'textbook', # This might need to be more dynamic if you upload other types
                    'is_processed': is_processed_flag,
                    'processed_dir_name': item_name # Store the directory name itself
                })
            else:
                # If no primary file found, but it's a _dir, maybe list it as an "unknown" processed item
                print(f"Directory {item_name} found but no primary file (.pdf, .docx, .txt) within it.")


    # Synchronize with mock_files_metadata (optional, if you're still using it for non-disk items)
    # For simplicity, this example now primarily relies on disk scan for _dir items.
    # If you have other types of files not in _dir structures, you'd merge them here.
    global mock_files_metadata
    mock_files_metadata[:] = files_on_disk # Replace mock data with fresh scan
    # save_mock_data() # If you want to persist this scanned list

    return jsonify(files_on_disk)


# 3. 上传文件
@app.route('/api/upload', methods=['POST'])
@cross_origin()
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request. Make sure your form field name is "file".'}), 400

    file = request.files['file']
    file_type_tag = request.form.get('file_type_tag', 'general_document') # More generic default

    if file.filename == '':
        return jsonify({'error': 'No selected file or empty filename.'}), 400

    original_filename = secure_filename(file.filename) # Sanitize original filename
    name_without_ext, ext = os.path.splitext(original_filename)

    # Generate a unique base name for processing (e.g., "Essay", "Essay_1")
    # This base name is used by some of your sub-scripts.
    textbook_base_name_unique = name_without_ext
    if not textbook_base_name_unique: # Handle cases like ".pdf"
        textbook_base_name_unique = f"file_{int(time.time() * 1000)}"
    
    # Ensure textbook_base_name_unique is truly unique if multiple files with same name_without_ext are uploaded
    # This loop is for the *base name* before adding _dir.
    # The directory name uniqueness is handled separately.
    # This part might be redundant if the directory naming ensures overall uniqueness.
    # For simplicity, let's assume textbook_base_name_unique from filename is sufficient for script args,
    # and the directory name `processed_dir_full_name` is the primary unique identifier.

    # Generate a unique directory name for this upload instance (e.g., "Essay_dir", "Essay_1_dir")
    # This directory will store the original file and all its processed outputs.
    # This `processed_dir_full_name` will be the 'id' for the file in the frontend.
    base_for_dir = name_without_ext if name_without_ext else f"upload_{int(time.time())}"
    processed_dir_full_name = f"{base_for_dir}_dir"
    counter = 1
    while (Path(UPLOAD_FOLDER) / processed_dir_full_name).exists():
        processed_dir_full_name = f"{base_for_dir}_{counter}_dir"
        counter += 1
    
    # The actual base name passed to scripts (without _dir)
    # If `base_for_dir` was modified by counter, reflect that here for consistency if scripts rely on it.
    # Example: if original was "Essay.pdf", base_for_dir is "Essay".
    # If "Essay_dir" exists, next is "Essay_1_dir", so script base name could be "Essay_1".
    # This ensures scripts get a unique base name if they create files based on it directly.
    script_arg_base_name = base_for_dir if counter == 1 else f"{base_for_dir}_{counter-1}"


    target_processing_directory = Path(UPLOAD_FOLDER) / processed_dir_full_name
    target_processing_directory.mkdir(parents=True, exist_ok=True)

    # Save the original uploaded file into this unique directory
    path_to_saved_original_file = target_processing_directory / original_filename
    
    try:
        file.save(str(path_to_saved_original_file))
        print(f"Original file '{original_filename}' saved to: {path_to_saved_original_file}")

        file_ext_type = ext.replace('.', '').lower() if ext else 'unknown'
        new_file_meta = {
            'id': processed_dir_full_name, # The unique directory is the ID
            'name': original_filename,      # Display name is the original filename
            'size': f"{round(path_to_saved_original_file.stat().st_size / (1024 * 1024), 2)}MB",
            'type': file_ext_type,
            'uploadDate': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'file_type_tag': file_type_tag,
            'is_processed': False, # Initially not processed
            'processed_dir_name': processed_dir_full_name # Redundant with 'id' but clear
        }

        # Add to in-memory list (if you're still using mock_files_metadata for immediate UI update)
        # mock_files_metadata.append(new_file_meta)
        # save_mock_data() # If persisting this list

        print(
            f"Launching background processing for base name '{script_arg_base_name}' (Original: '{original_filename}', Target Dir: '{processed_dir_full_name}')...")
        
        thread = threading.Thread(target=run_processing_pipeline_in_thread,
                                  args=(script_arg_base_name, original_filename, processed_dir_full_name))
        thread.start()

        return jsonify({
            'message': 'File uploaded successfully. Processing started in background.',
            'fileId': new_file_meta['id'], # Return the unique directory name as ID
            'fileName': new_file_meta['name'],
            'fileTypeTag': file_type_tag,
            'is_processed': False
        }), 200
    except Exception as e:
        print(f"Error saving file or starting processing: {e}")
        # Clean up created directory if save failed mid-way?
        if target_processing_directory.exists():
            shutil.rmtree(target_processing_directory) # Example cleanup
        return jsonify({'error': f'Failed to upload file or start processing: {str(e)}'}), 500


# 4. 获取对话历史列表
@app.route('/api/chat-history', methods=['GET'])
@cross_origin()
def get_chat_history():
    # This part remains largely the same, assuming chat history is linked by 'file_id'
    # which should correspond to 'processed_dir_name'
    history_with_files = []
    # Ensure mock_files_metadata is up-to-date if you rely on it here
    # Or fetch file metadata directly based on IDs stored in chat_entry
    current_files_map = {f['id']: f for f in mock_files_metadata} # Assuming mock_files_metadata is current

    for chat_entry in mock_chat_history:
        entry_copy = chat_entry.copy()
        entry_copy['related_files_meta'] = []
        for file_id_in_chat in entry_copy.get('related_file_ids', []):
            # file_id_in_chat should be a processed_dir_name like "Essay_dir"
            file_meta = current_files_map.get(file_id_in_chat)
            if file_meta:
                entry_copy['related_files_meta'].append(file_meta)
        history_with_files.append(entry_copy)
    return jsonify(history_with_files)


# 5. 获取某个对话的消息
@app.route('/api/chat/<chat_id>/messages', methods=['GET'])
@cross_origin()
def get_chat_messages(chat_id):
    messages = mock_chat_messages.get(chat_id, [])
    related_files_meta = []
    chat_entry = next((c for c in mock_chat_history if c['id'] == chat_id), None)

    if chat_entry and chat_entry.get('related_file_ids'):
        current_files_map = {f['id']: f for f in mock_files_metadata} # Get current files
        for file_id in chat_entry['related_file_ids']:
            file_meta = current_files_map.get(file_id)
            if file_meta:
                related_files_meta.append(file_meta)
    return jsonify({"messages": messages, "related_files_meta": related_files_meta})


# 6. 发送新消息到大模型 (RAG Chat)
@app.route('/api/chat', methods=['POST'])
@cross_origin()
def send_chat_message():
    data = request.get_json()
    user_message_content = data.get('message')
    chat_id = data.get('chatId') # Existing chat or null for new
    model_settings = data.get('modelSettings', {})
    # relatedFileIds are expected to be the 'processed_dir_name' values (e.g., "Essay_dir")
    related_file_ids_from_request = data.get('relatedFileIds', [])


    if not user_message_content:
        return jsonify({'error': 'Message content cannot be empty'}), 400

    new_chat_id = chat_id
    current_chat_entry = None

    # Ensure mock_files_metadata is current to resolve file details
    # This could be done by calling a function that refreshes it or ensuring get_files() was called recently by frontend
    # For simplicity, assume mock_files_metadata is reasonably up-to-date or use a direct lookup if needed.
    
    current_files_on_server_map = {f['id']: f for f in mock_files_metadata} # Using the global list

    if not new_chat_id: # New chat
        new_chat_id = f'c{int(time.time() * 1000)}'
        new_chat_title = f"新对话 {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        current_chat_entry = {
            'id': new_chat_id,
            'title': new_chat_title,
            'lastActive': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'related_file_ids': related_file_ids_from_request # Store the dir names
        }
        mock_chat_history.insert(0, current_chat_entry)
        mock_chat_messages[new_chat_id] = []
    else: # Existing chat
        for chat_entry_obj in mock_chat_history:
            if chat_entry_obj['id'] == new_chat_id:
                current_chat_entry = chat_entry_obj
                # Update related files if they've changed for this message in an existing chat
                current_chat_entry['related_file_ids'] = related_file_ids_from_request
                current_chat_entry['lastActive'] = datetime.now().strftime('%Y-%m-%d %H:%M')
                break
    
    current_chat_msgs_list = mock_chat_messages.get(new_chat_id, [])
    current_chat_msgs_list.append({
        'id': f'm_user_{int(time.time() * 1000)}',
        'sender': 'user',
        'content': user_message_content,
        'timestamp': datetime.now().strftime('%H:%M')
    })
    # mock_chat_messages[new_chat_id] = current_chat_msgs_list # Already a reference, modification is enough

    final_ai_response_json = {"text": "抱歉，AI助手暂时无法响应。请确保您的API Key正确且网络畅通。", "citations": []}

    try:
        messages_for_llm = []
        context_str = ""
        citations_for_response = []
        all_knowledge_points = [] # This will store dicts like {"id": "...", "text": "...", "doc_name": "...", "chapter_name": "..."}

        if related_file_ids_from_request:
            print(f"DEBUG RAG: Processing relatedFileIds: {related_file_ids_from_request}")
            for processed_dir_id in related_file_ids_from_request:
                # processed_dir_id is like "Essay_dir"
                file_meta_for_rag = current_files_on_server_map.get(processed_dir_id)

                if not file_meta_for_rag:
                    print(f"WARNING RAG: Metadata not found for file ID '{processed_dir_id}'. Skipping for RAG.")
                    continue
                
                original_doc_name_for_citation = file_meta_for_rag.get('name', processed_dir_id) # e.g. "Essay.pdf"

                # Path to the processed directory for this file ID
                # e.g., uploads/Essay_dir/
                current_processed_dir_path = Path(UPLOAD_FOLDER) / processed_dir_id
                
                # Path to the catalog_with_segments.json within that directory
                # e.g., uploads/Essay_dir/textbook_information/catalog_with_segments.json
                catalog_file_path = current_processed_dir_path / "textbook_information" / "catalog_with_segments.json"

                if file_meta_for_rag.get('is_processed') and catalog_file_path.exists():
                    print(f"DEBUG RAG: Loading knowledge points from processed file: {original_doc_name_for_citation} (catalog: {catalog_file_path})")
                    try:
                        with open(catalog_file_path, 'r', encoding='utf-8') as f:
                            catalog_data = json.load(f)
                        
                        # Recursive function to extract KPs from catalog structure
                        def collect_kps_from_catalog_nodes(nodes, current_chapter_path_str=""):
                            kps_list = []
                            for node_item in nodes:
                                node_id_part = node_item.get('generated_path_id', node_item.get('id', 'unknown_node'))
                                full_node_path = f"{current_chapter_path_str}.{node_id_part}" if current_chapter_path_str else node_id_part
                                
                                if node_item.get('knowledge_points') and isinstance(node_item['knowledge_points'], list):
                                    for idx, kp_text_content in enumerate(node_item['knowledge_points']):
                                        kps_list.append({
                                            "id": f"{full_node_path}_kp_{idx}", # Unique ID for the KP
                                            "text": kp_text_content,
                                            "doc_name": original_doc_name_for_citation,
                                            "chapter_name": node_item.get('name', 'Unknown Chapter') # Name of the current catalog node
                                        })
                                if node_item.get('children') and isinstance(node_item['children'], list):
                                    kps_list.extend(collect_kps_from_catalog_nodes(node_item['children'], full_node_path))
                            return kps_list

                        # Assuming catalog_data might be a dict with a 'chapters' key, or a direct list of chapters
                        chapters_list_from_catalog = []
                        if isinstance(catalog_data, dict) and 'chapters' in catalog_data and isinstance(catalog_data['chapters'], list):
                             chapters_list_from_catalog = catalog_data['chapters']
                        elif isinstance(catalog_data, list): # If the JSON root is the list of chapters
                             chapters_list_from_catalog = catalog_data
                        
                        if chapters_list_from_catalog:
                            kps_from_this_file = collect_kps_from_catalog_nodes(chapters_list_from_catalog)
                            all_knowledge_points.extend(kps_from_this_file)
                            print(f"DEBUG RAG: Loaded {len(kps_from_this_file)} KPs from {original_doc_name_for_citation}.")
                        else:
                            print(f"WARNING RAG: 'chapters' array not found or empty in catalog for {original_doc_name_for_citation}.")

                    except Exception as e_catalog:
                        print(f"ERROR RAG: Failed to load/parse catalog {catalog_file_path} for {original_doc_name_for_citation}: {e_catalog}")
                else:
                    # Fallback: File not processed or catalog missing, use raw text extraction
                    print(f"WARNING RAG: File '{original_doc_name_for_citation}' (ID: {processed_dir_id}) not marked as processed or catalog missing. Falling back to raw text chunking.")
                    
                    # The original file is stored directly in the processed_dir_id
                    # e.g., uploads/Essay_dir/Essay.pdf
                    raw_file_path_for_extraction = current_processed_dir_path / original_doc_name_for_citation
                    
                    if raw_file_path_for_extraction.exists():
                        # extract_text_from_file returns (list_of_chunks, error_message)
                        # Each chunk is a dict: {"id": "chunk_X", "text": "...", "original_full_text": "..."}
                        raw_text_chunks, extract_err = extract_text_from_file(raw_file_path_for_extraction, file_meta_for_rag.get('type', 'unknown'))
                        if extract_err:
                             print(f"WARNING RAG: Error extracting raw text from {original_doc_name_for_citation}: {extract_err}")
                        if raw_text_chunks:
                            for r_chunk in raw_text_chunks:
                                all_knowledge_points.append({
                                    "id": f"{processed_dir_id}_raw_{r_chunk['id']}",
                                    "text": r_chunk['text'],
                                    "doc_name": original_doc_name_for_citation,
                                    "chapter_name": "Raw Text Chunk (Unprocessed File)"
                                })
                            print(f"DEBUG RAG: Added {len(raw_text_chunks)} raw text chunks from {original_doc_name_for_citation}.")
                    else:
                        print(f"WARNING RAG: Original file {raw_file_path_for_extraction} not found for raw text extraction.")
        
        print(f"DEBUG RAG: Total knowledge points collected for RAG: {len(all_knowledge_points)}")
        # (The rest of your RAG logic: jieba scoring, selecting top K, building prompt)
        # This part is complex and assumed to be mostly correct from your original code.
        # Key is that `all_knowledge_points` is now populated correctly.

        relevant_chunks_for_llm = []
        if all_knowledge_points:
            # Simplified relevance scoring for brevity (your jieba logic is more sophisticated)
            # This is where your existing jieba-based scoring and selection would go.
            # For this example, let's just take the first few if they exist.
            # Replace this with your actual scoring and selection logic.
            
            query_segmented_words = jieba.lcut(user_message_content.lower())
            stop_words = {"的", "了", "是", "在", "我", "你", "他", "她", "它", "分析", "文件", "请", "如何", "什么",
                          "这个", "那个", "根据", "文档", "回答", "问题", "并", "和", "或", "等", "以", "从", "中",
                          "于", "对", "为", "所", "其", "则", "将", "与", "关于", "有", "也", "都", "还", "是什么",
                          "的", "地", "得"}
            query_words = {word for word in query_segmented_words if len(word) > 1 and word not in stop_words}
            print(f"DEBUG RAG: Processed Query Words (jieba): {query_words}")

            scored_chunks = []
            for kp_item in all_knowledge_points:
                if not kp_item.get('text'): continue
                kp_segmented_words = set(jieba.lcut(kp_item['text'].lower()))
                score = sum(1 for word in query_words if word in kp_segmented_words)
                if score > 0:
                    scored_chunks.append((score, kp_item))
            
            scored_chunks.sort(key=lambda x: x[0], reverse=True)
            
            MAX_CONTEXT_CHUNKS = 8 # As in your original code
            citation_id_counter = 1
            seen_kp_identifiers_for_rag = set()

            for score, kp_data in scored_chunks:
                if citation_id_counter > MAX_CONTEXT_CHUNKS: break
                
                # Create a unique identifier for this KP instance to avoid duplicates in context
                kp_unique_id_for_context = f"{kp_data['doc_name']}_{kp_data.get('chapter_name', 'N/A')}_{kp_data['text'][:50]}"
                if kp_unique_id_for_context in seen_kp_identifiers_for_rag: continue
                
                relevant_chunks_for_llm.append({
                    "id": f"[{citation_id_counter}]", # Citation marker for LLM
                    "text": kp_data['text'],
                    "original_doc_name": kp_data['doc_name'],
                    "chapter_name": kp_data.get('chapter_name', 'Unknown Chapter') # For better context in prompt
                })
                citations_for_response.append({ # For frontend display
                    "id": str(citation_id_counter),
                    "doc_name": kp_data['doc_name'],
                    "text": kp_data['text'] # Full text for citation popover
                })
                seen_kp_identifiers_for_rag.add(kp_unique_id_for_context)
                citation_id_counter += 1
            print(f"DEBUG RAG: Final {len(relevant_chunks_for_llm)} relevant KPs selected for LLM context.")


        # Constructing the prompt for LLM
        system_instruction_prefix = "你是一个严谨、专业的AI助手，擅长分析文档并提供清晰、准确、且带有引用的回答。你必须严格依据提供的参考资料进行回答，不允许虚构或依赖你的预训练知识回答参考资料中没有的信息。"
        user_prompt_content = f"用户问题：{user_message_content}"

        if related_file_ids_from_request and relevant_chunks_for_llm:
            context_str += "以下是一些用户选择的参考资料。请注意，你的回答必须完全基于这些资料，如果答案不在其中，请明确说明。\n\n"
            for chunk_info_for_llm in relevant_chunks_for_llm:
                context_str += f"### 参考资料：文档 '{chunk_info_for_llm['original_doc_name']}'，片段 {chunk_info_for_llm['id']}：来自章节 '{chunk_info_for_llm.get('chapter_name', '未知章节')}'\n"
                context_str += "```text\n"
                context_str += chunk_info_for_llm['text'] + "\n"
                context_str += "```\n\n"
            context_str += "请根据上述提供的参考资料，准确、简洁地回答以下用户的问题。**在回答中，如果使用了参考资料，务必在相关内容后用方括号加数字的形式引用，例如：某个概念的定义[1]，某个论点[2]。** 如果参考资料中没有提及相关信息，请直接回答‘抱歉，根据我当前掌握的资料，无法从提供的参考资料中找到此问题的答案。’\n\n"
        elif related_file_ids_from_request and not relevant_chunks_for_llm:
            context_str = "用户提问，但无法从选择的文件中找到与问题相关的参考资料。我将完全根据我已有的知识进行回答。如果您的原始问题是关于特定文档的，请确保文档内容中包含您的问题。请注意，我无法访问您本地的文件。"
        else: # No files selected
            context_str = "用户提问，未选择任何参考资料。我将完全根据我已有的知识进行回答。\n\n"

        messages_for_llm.append({"role": "system", "content": system_instruction_prefix + context_str})
        messages_for_llm.append({"role": "user", "content": user_prompt_content})

        selected_model_for_llm = model_settings.get('model', DEFAULT_DEEPSEEK_MODEL)
        # Ensure model is valid, etc.
        
        print(f"Calling DeepSeek LLM with model: {selected_model_for_llm}, temperature: {model_settings.get('temperature', 0.7)}")
        # print(f"Prompt sent to LLM (first 500 chars of context): {messages_for_llm[0]['content'][:500]}...")
        # print(f"User part of prompt: {messages_for_llm[1]['content']}")

        completion = openai_client_for_chat_rag.chat.completions.create(
            model=selected_model_for_llm,
            messages=messages_for_llm,
            temperature=model_settings.get('temperature', 0.7),
            timeout=30.0 # Consider making this configurable
        )
        ai_response_content = completion.choices[0].message.content
        print(f"AI Response (first 200 chars): {ai_response_content[:200]}...")

        final_ai_response_json["text"] = ai_response_content
        final_ai_response_json["citations"] = citations_for_response # Citations prepared earlier

    # (Your existing error handling for APIConnectionError, RateLimitError, etc.)
    except APIConnectionError as e:
        print(f"ERROR: DeepSeek API Connection Error: {e}")
        final_ai_response_json["text"] = f"无法连接到DeepSeek服务，请检查网络或代理设置: {str(e)}"
        return jsonify({'aiResponse': final_ai_response_json}), 503 # Service Unavailable
    except RateLimitError as e:
        print(f"ERROR: DeepSeek API Rate Limit Exceeded: {e}")
        final_ai_response_json["text"] = f"请求DeepSeek太频繁，请稍后再试: {str(e)}"
        return jsonify({'aiResponse': final_ai_response_json}), 429 # Too Many Requests
    except APIStatusError as e:
        print(f"ERROR: DeepSeek API returned non-200 status: {e.status_code} - {e.response.json().get('message', '未知错误') if e.response else 'No response data'}")
        error_message_from_api = "未知API错误"
        if e.response:
            try:
                error_message_from_api = e.response.json().get('message', '未知API错误')
            except json.JSONDecodeError:
                error_message_from_api = e.response.text # If not JSON
        
        if e.status_code == 401: # Unauthorized
            final_ai_response_json["text"] = f"DeepSeek服务认证失败，请检查API Key是否有效或账户余额。"
        else:
            final_ai_response_json["text"] = f"DeepSeek服务内部错误: {e.status_code} - {error_message_from_api}"
        return jsonify({'aiResponse': final_ai_response_json}), e.status_code if e.status_code else 500
    except Exception as e:
        print(f"ERROR: General error calling DeepSeek API or during RAG processing: {e}")
        import traceback
        traceback.print_exc() # Print full traceback for general errors
        final_ai_response_json["text"] = f"抱歉，DeepSeek助手在处理您的请求时遇到问题: {str(e)}"
        return jsonify({'aiResponse': final_ai_response_json}), 500 # Internal Server Error


    current_chat_msgs_list.append({
        'id': f'm_ai_{int(time.time() * 1000) + 1}',
        'sender': 'ai',
        'content': final_ai_response_json['text'],
        'timestamp': datetime.now().strftime('%H:%M'),
        'citations_data': final_ai_response_json['citations']
    })
    # mock_chat_messages[new_chat_id] = current_chat_msgs_list # Already modified by reference

    # Update chat history last active time (if current_chat_entry was found/created)
    # This is already handled above when current_chat_entry is set or updated.

    save_mock_data() # Persist chat messages and history

    response_data_to_frontend = {
        'aiResponse': final_ai_response_json,
        'chatId': new_chat_id # Always return the chatId
    }
    if not chat_id: # If it was a new chat, also send newChatId and newChatTitle
        response_data_to_frontend['newChatId'] = new_chat_id
        response_data_to_frontend['newChatTitle'] = current_chat_entry['title'] if current_chat_entry else "新对话"
        
    return jsonify(response_data_to_frontend), 200


# 7. 导出对话为 DOCX 文件
@app.route('/api/export-chat/<chat_id>', methods=['GET'])
@cross_origin()
def export_chat(chat_id):
    messages = mock_chat_messages.get(chat_id, [])
    if not messages:
        return jsonify({'error': 'No messages found for this chat ID.'}), 404

    chat_title = "Chat History"
    chat_entry = next((c for c in mock_chat_history if c['id'] == chat_id), None)
    if chat_entry:
        chat_title = chat_entry.get('title', f"Chat-{chat_id}")

    safe_chat_title = secure_filename(chat_title)
    if not safe_chat_title: # Handle empty or all-special-char titles
        safe_chat_title = f"chat_export_{int(time.time())}"

    document = docx.Document()
    document.add_heading(f"对话历史 - {chat_title}", level=1)

    for msg in messages:
        sender = "您" if msg['sender'] == 'user' else "AI助手"
        timestamp = msg['timestamp']
        p = document.add_paragraph()
        p.add_run(f"[{timestamp}] {sender}: ").bold = True
        p.add_run(msg['content'])

        if msg['sender'] == 'ai' and msg.get('citations_data'):
            document.add_paragraph("--- 引用 ---", style='Intense Quote') # Example style
            for citation in msg['citations_data']:
                document.add_paragraph(f"[{citation['id']}] (来自 {citation['doc_name']}): {citation['text']}", style='List Paragraph')
            document.add_paragraph("-----------") # Separator

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{safe_chat_title}_对话历史.docx'
    )

# --- Endpoint to get the main textbook orgchart (from get_orgchart.py) ---
@app.route('/api/files/<string:processed_dir_name>/orgchart', methods=['GET'])
@cross_origin()
def get_orgchart_json(processed_dir_name):
    print(f"\n--- Request for main textbook orgchart.json for directory: {processed_dir_name} ---")
    # The main orgchart is usually associated with the textbook_base_name, not necessarily the full processed_dir_name
    # if processed_dir_name can include suffixes like _Exam1_dir.
    # Assuming processed_dir_name IS the directory containing textbook_information/textbook_orgchart.json
    # For example, if get_orgchart.py saves its output to uploads/MyBook_dir/textbook_information/textbook_orgchart.json
    # then processed_dir_name should be "MyBook_dir".

    # If get_orgchart.py saves based on textbook_base_name (e.g. "MyBook"), then:
    # textbook_base_name = processed_dir_name.replace("_dir", "") # Simplistic, adjust if needed
    # orgchart_file_path = Path(UPLOAD_FOLDER) / textbook_base_name / "textbook_information" / "textbook_orgchart.json"
    
    # Current assumption: processed_dir_name is the direct parent of "textbook_information" for this file
    orgchart_file_path = Path(UPLOAD_FOLDER) / processed_dir_name / "textbook_information" / "textbook_orgchart.json"
    
    print(f"Attempting to read main textbook orgchart from: {orgchart_file_path}")

    if not orgchart_file_path.is_file(): # Check if it's a file specifically
        print(f"Error: Main textbook orgchart.json not found or not a file at {orgchart_file_path}")
        file_meta = next((f for f in mock_files_metadata if f.get('id') == processed_dir_name), None)
        display_name = file_meta['name'] if file_meta else processed_dir_name
        return jsonify({
            'error': f"Main organizational chart data not found for '{display_name}'.",
            'detail': f"Expected file not found at server path: .../{processed_dir_name}/textbook_information/textbook_orgchart.json"
        }), 404

    try:
        with open(orgchart_file_path, 'r', encoding='utf-8') as f:
            orgchart_data = json.load(f)
        print(f"Successfully loaded main textbook orgchart.json for {processed_dir_name}")
        return jsonify(orgchart_data), 200
    except json.JSONDecodeError as e:
        print(f"Error: Failed to decode JSON from {orgchart_file_path}: {e}")
        return jsonify({'error': 'Failed to parse main org chart data. JSON may be corrupted.', 'detail': str(e)}), 500
    except Exception as e:
        print(f"Error: An unexpected error occurred while reading {orgchart_file_path}: {e}")
        return jsonify({'error': 'Unexpected server error for main org chart.', 'detail': str(e)}), 500


# --- NEW ENDPOINT: Get Questions-Filtered Org Chart JSON ---
@app.route('/api/files/<string:processed_dir_name>/questions-orgchart', methods=['GET'])
@cross_origin()
def get_questions_filtered_orgchart_json(processed_dir_name):
    """
    Serves the questions_filtered_textbook_orgchart.json file for a given processed book directory.
    'processed_dir_name' is expected to be the unique directory name like 'MyBook_Exam1_dir',
    which was used when running get_questions_orgchart.py.
    """
    print(f"\n--- Request for questions-filtered-orgchart.json for directory: {processed_dir_name} ---")

    # The file is generated by get_questions_orgchart.py and stored in:
    # uploads/<processed_dir_name>/textbook_information/questions_filtered_textbook_orgchart.json
    target_json_filename = "questions_filtered_textbook_orgchart.json"
    questions_orgchart_file_path = Path(UPLOAD_FOLDER) / processed_dir_name / "textbook_information" / target_json_filename
    
    print(f"Attempting to read questions-filtered orgchart from: {questions_orgchart_file_path}")

    if not questions_orgchart_file_path.is_file(): # Ensure it's a file
        print(f"Error: Questions-filtered orgchart JSON not found or not a file at {questions_orgchart_file_path}")
        # Try to get a display name for the error message
        file_meta = next((f for f in mock_files_metadata if f.get('id') == processed_dir_name), None)
        display_name = file_meta['name'] if file_meta else processed_dir_name
        return jsonify({
            'error': f"Questions-specific organizational chart data not found for '{display_name}'.",
            'detail': f"Expected file '{target_json_filename}' not found in .../{processed_dir_name}/textbook_information/"
        }), 404

    try:
        with open(questions_orgchart_file_path, 'r', encoding='utf-8') as f:
            questions_orgchart_data = json.load(f)
        print(f"Successfully loaded {target_json_filename} for {processed_dir_name}")
        return jsonify(questions_orgchart_data), 200
    except json.JSONDecodeError as e:
        print(f"Error: Failed to decode JSON from {questions_orgchart_file_path}: {e}")
        return jsonify({'error': 'Failed to parse questions org chart data. The JSON file may be corrupted.', 'detail': str(e)}), 500
    except Exception as e:
        print(f"Error: An unexpected error occurred while reading {questions_orgchart_file_path}: {e}")
        return jsonify({'error': 'An unexpected server error occurred while retrieving the questions org chart.', 'detail': str(e)}), 500


# --- 新功能：答题方法和模板提取相关路由 ---
@app.route('/api/templates/list', methods=['GET'])
@cross_origin()
def list_templates():
    templates_methods = template_manager.get_all_templates_methods()
    return jsonify(templates_methods)


@app.route('/api/templates/extract-from-file', methods=['POST'])
@cross_origin()
def extract_templates_from_file():
    data = request.get_json()
    file_id = data.get('fileId') # This file_id is the processed_dir_name

    if not file_id:
        return jsonify({"error": "File ID (processed_dir_name) is required."}), 400

    # Find the file metadata using the file_id (which is processed_dir_name)
    # Assuming mock_files_metadata is up-to-date or use get_files() logic to find it.
    file_meta = next((f for f in mock_files_metadata if f['id'] == file_id), None)
    if not file_meta:
        # Attempt to scan disk if not in mock_files_metadata (e.g., if mock data is not persisted/reloaded)
        # This is a simplified fallback; ideally, file metadata should be consistently available.
        temp_item_path = Path(UPLOAD_FOLDER) / file_id
        if temp_item_path.is_dir():
             # Try to find an original file inside to get its name and type
            original_file_in_dir = None
            for ext_type in ['.pdf', '.docx', '.txt']: # common types
                for f_content in temp_item_path.iterdir():
                    if f_content.is_file() and f_content.suffix.lower() == ext_type:
                        original_file_in_dir = f_content
                        break
                if original_file_in_dir: break
            
            if original_file_in_dir:
                file_meta = {
                    'id': file_id, 
                    'name': original_file_in_dir.name, 
                    'type': original_file_in_dir.suffix.lower().replace('.', ''),
                    'file_type_tag': 'textbook' # default
                }
            else: # Still couldn't determine original file details
                 file_meta = {'id': file_id, 'name': file_id, 'type': 'unknown', 'file_type_tag': 'unknown'}
        else:
            return jsonify({"error": f"File metadata not found for ID: {file_id}"}), 404


    # Path to the original file within its processed directory
    # e.g., uploads/Essay_dir/Essay.pdf
    path_to_original_in_processed_dir = Path(UPLOAD_FOLDER) / file_meta['id'] / file_meta['name']

    if not path_to_original_in_processed_dir.is_file():
        return jsonify(
            {"error": f"Original file '{file_meta['name']}' not found on disk at: {path_to_original_in_processed_dir}"}), 404

    # extract_text_from_file returns (chunks_list, error_message_str)
    full_text_chunks, error_msg_extraction = extract_text_from_file(path_to_original_in_processed_dir, file_meta['type'])
    
    if error_msg_extraction or not full_text_chunks:
        # Note: extract_text_from_file returns error_msg as string, not Exception object 'e'
        actual_error_detail = error_msg_extraction if error_msg_extraction else "No text chunks extracted."
        print(f"DEBUG APP: Failed to extract text chunks from file: {actual_error_detail}")
        return jsonify({"error": f"Failed to extract text from file: {actual_error_detail}"}), 500

    full_text_content = "\n\n".join([c['text'] for c in full_text_chunks])
    file_type_display = file_meta.get('file_type_tag', '未知类型')

    qa_pairs_from_llm = []
    try:
        qa_pairs_from_llm = llm_client_for_tasks.extract_qa_pairs_from_document(
            full_text_content, document_type_tag=file_type_display
        )
        print(f"DEBUG APP: LLM extracted Q&A pairs (count: {len(qa_pairs_from_llm)}): {str(qa_pairs_from_llm[:min(3, len(qa_pairs_from_llm))])[:200]}...") # Limit log size

        if not isinstance(qa_pairs_from_llm, list):
            print(f"DEBUG APP: LLM extract_qa_pairs_from_document returned non-list: {qa_pairs_from_llm}")
            return jsonify(
                {"error": f"LLM failed to extract valid Q&A pairs (returned non-list)."}), 500 # Avoid echoing potentially large LLM output
    except Exception as e_llm_extract:
        print(f"ERROR APP: Error calling LLM for Q&A pair extraction: {e_llm_extract}")
        return jsonify({"error": f"Failed to extract Q&A pairs from document: {str(e_llm_extract)}"}), 500

    if not qa_pairs_from_llm: # Handles both empty list and if LLM returned something invalid that became empty
        print(f"DEBUG APP: No valid Q&A pairs identified by LLM for file {file_meta['name']}.")
        return jsonify({"message": "Successfully analyzed file, but no valid Q&A pairs were identified.",
                        "extracted_data": []}), 200

    added_count = 0
    extracted_templates_methods = []

    for i, qa_pair_item in enumerate(qa_pairs_from_llm):
        question = qa_pair_item.get("question")
        answer = qa_pair_item.get("answer")

        if question and answer:
            print(f"DEBUG APP: Processing Q&A pair {i + 1}: Q='{str(question)[:50]}...', A='{str(answer)[:50]}...'")
            try:
                # Assuming question_rewriter_module.extract_and_save_qa_template_method exists and works as intended
                extract_result = question_rewriter_module.extract_and_save_qa_template_method(question, answer)
                print(f"DEBUG APP: Result from question_rewriter_module for pair {i + 1}: {extract_result}")

                if "error" not in extract_result: # Assuming success if no 'error' key
                    q_template = extract_result.get("question_template")
                    a_method = extract_result.get("answer_method")
                    if q_template and a_method:
                        if template_manager.add_template_method(q_template, a_method): # Assumes this returns True on success/add
                            added_count += 1
                            extracted_templates_methods.append({
                                "question_template": q_template,
                                "answer_method": a_method
                            })
                            print(f"DEBUG APP: Successfully added template/method for pair {i + 1}.")
                        # else: already exists or failed to add, template_manager should handle logging
                    # else: incomplete data from LLM for template extraction
                # else: error reported by question_rewriter_module
            except Exception as e_template_extract:
                print(f"ERROR APP: Error processing QA pair {i + 1} for template extraction: {e_template_extract}")
        # else: incomplete Q&A pair from initial LLM extraction

    return jsonify({"message": f"Successfully extracted and saved {added_count} templates and methods from file.",
                    "extracted_data": extracted_templates_methods}), 200


@app.route('/api/templates/rewrite-answer', methods=['POST'])
@cross_origin()
def rewrite_answer_api():
    data = request.get_json()
    question = data.get('question')
    original_answer = data.get('originalAnswer')
    method_index = data.get('methodIndex') # Expecting an index or ID for the method

    if not all([question, original_answer, method_index is not None]): # Check method_index presence
        return jsonify({"error": "Missing question, original answer, or method identifier."}), 400

    try:
        # Assuming rewrite_answer_with_selected_method handles the method_index correctly
        rewritten_answer = question_rewriter_module.rewrite_answer_with_selected_method(
            question, original_answer, method_index
        )
        if isinstance(rewritten_answer, str) and "Error:" in rewritten_answer: # Check if it returns error string
            return jsonify({"error": rewritten_answer}), 500
        return jsonify({"rewrittenAnswer": rewritten_answer}), 200
    except Exception as e:
        print(f"Error in rewrite_answer_api: {e}")
        return jsonify({"error": f"Failed to rewrite answer: {str(e)}"}), 500


# --- 服务器启动 ---
if __name__ == '__main__':
    load_mock_data() # Load any persisted mock data on startup
    print("\n--- Flask URL Map ---")
    for rule in app.url_map.iter_rules():
        print(f"Endpoint: {rule.endpoint}, Methods: {rule.methods}, Rule: {rule.rule}")
    print("---------------------\n")

    print(f"\n--- Starting Flask backend server ---")
    print(f"DeepSeek API Key loaded: {'Yes' if deepseek_api_key else 'No'}")
    print(f"DashScope API Key loaded: {'Yes' if dashscope_api_key else 'No'}")
    print(f"DeepSeek Base URL: {DEEPSEEK_BASE_URL}")
    print(f"Uploads directory: {os.path.abspath(UPLOAD_FOLDER)}")
    print(f"Mock data file: {os.path.abspath(MOCK_DATA_FILE)}")
    print(f"Templates/Methods file: {os.path.abspath(TEMPLATES_FILE)}\n")
    
    # Added host='0.0.0.0' to make it accessible from network if needed,
    # debug=True is fine for development.
    app.run(host='0.0.0.0', debug=True, port=5000)
